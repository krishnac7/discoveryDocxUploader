import os,xlwt
from docx import Document

def splitDoc(documentDir,outputDir):

    
    book = xlwt.Workbook(encoding="utf-8")
    sheet1 = book.add_sheet("Sheet 1")
    headers=["Input_path","Output_path","Header","Text"]
    for i,header in enumerate(headers):
        sheet1.write(0,i,header)
    pointer = 1    

    if not os.path.exists(outputDir):
        os.makedirs(outputDir)

    files=os.listdir(documentDir)
    for file in files:
        if file.split(".")[1]=="docx" or file.split(".")[1]=="doc":
            print("[splitDoc] Processing {}".format(file))
            inputDocument = documentDir+"/"+file
            input_document = Document(inputDocument)
            paragraphs = input_document.paragraphs
            headings=[]

            for i,para in enumerate(paragraphs):
                for run in para.runs:
                    if run.bold and ":" in run.text:
                        headings.append(i)
                        
            for j,heading in enumerate(headings):
                document = Document()
                p = document.add_paragraph()
                run = p.add_run(paragraphs[heading].text)
                run.bold = True
                paras=""

                if not j+1>len(headings)-1:
                    for i in range(headings[j]+1,headings[j+1]):         
                        document.add_paragraph(paragraphs[i].text)
                        paras += "\n"+ paragraphs[i].text
                else:
                    for i in range(headings[j]+1,len(paragraphs)):
                        document.add_paragraph(paragraphs[i].text)
                        paras += "\n"+ paragraphs[i].text
                
                fileName = "{}{}_{}.docx".format(outputDir,file.split('.')[0],paragraphs[heading].text.split(":")[0])
                print("[splitDoc] Writing to {}".format(fileName))
                document.save(fileName)
                sheet1.write(pointer,0,inputDocument)
                sheet1.write(pointer,1,fileName)
                sheet1.write(pointer,2,paragraphs[heading].text.split(":")[0])
                sheet1.write(pointer,3,paras)
                pointer+=1
        print("\n\033[92m[splitDoc] Creating summary Excel....Complete\033[0m")
        book.save("{}summary.xls".format(outputDir))

